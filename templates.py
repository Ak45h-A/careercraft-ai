import streamlit as st
import sys, os
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

TEMPLATES = {
    "modern_purple": {
        "name": "Modern Purple",
        "emoji": "🟣",
        "desc": "Clean & professional with purple accents",
        "accent": "#6c63ff",
        "header_bg": "#6c63ff",
        "text": "#333",
        "font": "Segoe UI"
    },
    "minimal_black": {
        "name": "Minimal Black",
        "emoji": "⚫",
        "desc": "Ultra-minimal with bold typography",
        "accent": "#111",
        "header_bg": "#111",
        "text": "#222",
        "font": "Georgia"
    },
    "corporate_blue": {
        "name": "Corporate Blue",
        "emoji": "🔵",
        "desc": "Formal corporate style",
        "accent": "#1a56db",
        "header_bg": "#1a56db",
        "text": "#1e3a5f",
        "font": "Arial"
    },
    "creative_teal": {
        "name": "Creative Teal",
        "emoji": "🟢",
        "desc": "Modern creative industries look",
        "accent": "#0d9488",
        "header_bg": "#0d9488",
        "text": "#134e4a",
        "font": "Trebuchet MS"
    },
    "elegant_maroon": {
        "name": "Elegant Maroon",
        "emoji": "🔴",
        "desc": "Classic executive/senior look",
        "accent": "#9b1c1c",
        "header_bg": "#9b1c1c",
        "text": "#4c0519",
        "font": "Palatino"
    },
    "fresh_green": {
        "name": "Fresh Green",
        "emoji": "🌿",
        "desc": "Modern, youthful and energetic",
        "accent": "#16a34a",
        "header_bg": "#16a34a",
        "text": "#14532d",
        "font": "Verdana"
    },
}

def generate_html_resume(p: dict, template: dict, resume_text: str = "") -> str:
    name = p.get("name", "Your Name")
    email = p.get("email", "")
    phone = p.get("phone", "")
    city = p.get("city", "")
    linkedin = p.get("linkedin", "")
    github = p.get("github", "")
    summary = p.get("summary", "")
    skills = p.get("technical_skills", "")
    soft = p.get("soft_skills", "")
    certs = p.get("certifications", "")
    projects = p.get("projects", "")
    achievements = p.get("achievements", "")
    accent = template["accent"]
    header_bg = template["header_bg"]
    font = template["font"]

    exps_html = ""
    for exp in p.get("experiences", []):
        if exp.get("title"):
            exps_html += f"""
            <div class="entry">
                <div class="entry-header">
                    <strong>{exp.get('title','')}</strong> — {exp.get('company','')}
                    <span class="date">{exp.get('start','')} – {exp.get('end','')}</span>
                </div>
                <div class="entry-sub">{exp.get('location','')}</div>
                <div class="entry-desc">{exp.get('description','').replace(chr(10),'<br>')}</div>
            </div>"""

    edus_html = ""
    for edu in p.get("education", []):
        if edu.get("degree"):
            edus_html += f"""
            <div class="entry">
                <div class="entry-header">
                    <strong>{edu.get('degree','')}</strong>
                    <span class="date">{edu.get('year','')}</span>
                </div>
                <div class="entry-sub">{edu.get('institution','')} | {edu.get('grade','')}</div>
            </div>"""

    skills_html = ""
    if skills:
        for sk in skills.split(","):
            sk = sk.strip()
            if sk:
                skills_html += f'<span class="skill-badge">{sk}</span>'

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Resume - {name}</title>
<style>
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{
    font-family: '{font}', sans-serif;
    color: {template['text']};
    background: #f9f9f9;
    padding: 0;
  }}
  .page {{
    max-width: 900px;
    margin: 0 auto;
    background: white;
    box-shadow: 0 4px 20px rgba(0,0,0,0.1);
  }}
  .header {{
    background: {header_bg};
    color: white;
    padding: 40px 48px 32px;
  }}
  .header h1 {{
    font-size: 2.4rem;
    font-weight: 700;
    letter-spacing: 1px;
    margin-bottom: 6px;
  }}
  .header .tagline {{
    font-size: 1rem;
    opacity: 0.85;
    margin-bottom: 16px;
  }}
  .contact-row {{
    display: flex;
    flex-wrap: wrap;
    gap: 16px;
    font-size: 0.85rem;
    opacity: 0.9;
  }}
  .contact-row span {{ display: flex; align-items: center; gap: 4px; }}
  .body {{ display: flex; }}
  .sidebar {{
    width: 30%;
    background: #f5f5f7;
    padding: 28px 20px;
    min-height: 100%;
  }}
  .main {{ width: 70%; padding: 28px 32px; }}
  .section {{ margin-bottom: 24px; }}
  .section-title {{
    font-size: 0.75rem;
    font-weight: 700;
    text-transform: uppercase;
    letter-spacing: 2px;
    color: {accent};
    border-bottom: 2px solid {accent};
    padding-bottom: 4px;
    margin-bottom: 14px;
  }}
  .entry {{ margin-bottom: 16px; }}
  .entry-header {{
    display: flex;
    justify-content: space-between;
    align-items: flex-start;
    font-size: 0.95rem;
    font-weight: 600;
    color: #1a1a2e;
  }}
  .date {{
    font-size: 0.8rem;
    color: #888;
    font-weight: 400;
    white-space: nowrap;
    margin-left: 8px;
  }}
  .entry-sub {{ font-size: 0.82rem; color: #666; margin: 2px 0 6px; }}
  .entry-desc {{ font-size: 0.87rem; color: {template['text']}; line-height: 1.6; }}
  .skill-badge {{
    display: inline-block;
    background: {accent}22;
    color: {accent};
    border: 1px solid {accent}44;
    border-radius: 20px;
    padding: 3px 10px;
    font-size: 0.78rem;
    margin: 3px 2px;
    font-weight: 500;
  }}
  .summary-text {{
    font-size: 0.88rem;
    line-height: 1.7;
    color: #444;
  }}
  .soft-skills p {{
    font-size: 0.87rem;
    color: #555;
    margin-bottom: 4px;
    padding-left: 10px;
    border-left: 3px solid {accent};
    margin-bottom: 6px;
  }}
  @media print {{
    body {{ background: white; }}
    .page {{ box-shadow: none; }}
  }}
</style>
</head>
<body>
<div class="page">
  <div class="header">
    <h1>{name}</h1>
    <div class="tagline">{p.get('total_experience','') + ' Experience' if p.get('total_experience') else 'Professional Resume'}</div>
    <div class="contact-row">
      {'<span>📧 ' + email + '</span>' if email else ''}
      {'<span>📞 ' + phone + '</span>' if phone else ''}
      {'<span>📍 ' + city + '</span>' if city else ''}
      {'<span>🔗 ' + linkedin + '</span>' if linkedin else ''}
      {'<span>💻 ' + github + '</span>' if github else ''}
    </div>
  </div>

  <div class="body">
    <div class="sidebar">
      {'<div class="section"><div class="section-title">Skills</div>' + skills_html + '</div>' if skills_html else ''}

      {'<div class="section"><div class="section-title">Soft Skills</div><div class="soft-skills">' + ''.join(['<p>' + s.strip() + '</p>' for s in soft.split(',') if s.strip()]) + '</div></div>' if soft else ''}

      {'<div class="section"><div class="section-title">Certifications</div><div class="entry-desc">' + certs.replace(chr(10),'<br>') + '</div></div>' if certs else ''}

      {'<div class="section"><div class="section-title">Languages</div><div class="entry-desc">' + p.get('languages','') + '</div></div>' if p.get('languages') else ''}

      {'<div class="section"><div class="section-title">Interests</div><div class="entry-desc">' + p.get('hobbies','') + '</div></div>' if p.get('hobbies') else ''}
    </div>

    <div class="main">
      {'<div class="section"><div class="section-title">Professional Summary</div><p class="summary-text">' + summary + '</p></div>' if summary else ''}

      {'<div class="section"><div class="section-title">Work Experience</div>' + exps_html + '</div>' if exps_html else ''}

      {'<div class="section"><div class="section-title">Education</div>' + edus_html + '</div>' if edus_html else ''}

      {'<div class="section"><div class="section-title">Projects</div><div class="entry-desc">' + projects.replace(chr(10),'<br>') + '</div></div>' if projects else ''}

      {'<div class="section"><div class="section-title">Achievements</div><div class="entry-desc">' + achievements.replace(chr(10),'<br>') + '</div></div>' if achievements else ''}
    </div>
  </div>
</div>
</body>
</html>"""


def render():
    st.markdown('<h1>🎨 Visual Resume Templates</h1>', unsafe_allow_html=True)
    st.markdown('<p class="subtitle">Canva-style beautiful resume templates — pick, preview & download</p>', unsafe_allow_html=True)

    p = st.session_state.profile

    if not p.get("name"):
        st.markdown("""<div class="tip-box">⚠️ Please fill in <strong>My Profile</strong> first to preview templates with your data.</div>""", unsafe_allow_html=True)

    st.markdown("### Choose a Template")

    cols = st.columns(3)
    selected = st.session_state.get("selected_template", "modern_purple")

    for i, (key, tmpl) in enumerate(TEMPLATES.items()):
        with cols[i % 3]:
            is_sel = selected == key
            border = "border: 2px solid #6c63ff;" if is_sel else "border: 2px solid #2a2a3e;"
            st.markdown(f"""
            <div class="feature-card" style="{border}">
                <h3>{tmpl['emoji']} {tmpl['name']}</h3>
                <p>{tmpl['desc']}</p>
                <div style="margin-top:8px;">
                    <span style="display:inline-block; width:20px; height:20px; border-radius:50%;
                    background:{tmpl['accent']}; vertical-align:middle;"></span>
                    <span style="color:#9090aa; font-size:0.8rem; margin-left:6px;">{tmpl['accent']}</span>
                </div>
                {"<br><span class='badge'>✅ Selected</span>" if is_sel else ""}
            </div>
            """, unsafe_allow_html=True)
            if st.button(f"Select {tmpl['name']}", key=f"sel_{key}", use_container_width=True):
                st.session_state.selected_template = key
                st.rerun()

    st.markdown("<hr class='divider'>", unsafe_allow_html=True)

    selected_key = st.session_state.get("selected_template", "modern_purple")
    tmpl = TEMPLATES[selected_key]

    st.markdown(f"### Preview: {tmpl['emoji']} {tmpl['name']}")

    if p.get("name"):
        html = generate_html_resume(p, tmpl)

        col1, col2, col3 = st.columns(3)
        with col1:
            st.download_button("🌐 Download HTML", data=html, file_name=f"resume_{selected_key}.html",
                mime="text/html", use_container_width=True)
        with col2:
            st.download_button("📝 Download Text", data=st.session_state.get("generated_resume","No resume generated yet"),
                file_name="resume.txt", mime="text/plain", use_container_width=True)
        with col3:
            try:
                from docx import Document
                from io import BytesIO
                doc = Document()
                doc.add_heading(p.get("name", ""), 0)
                for exp in p.get("experiences", []):
                    if exp.get("title"):
                        doc.add_heading(f"{exp['title']} at {exp.get('company','')}", 2)
                        doc.add_paragraph(exp.get("description",""))
                buf = BytesIO()
                doc.save(buf)
                st.download_button("📄 Download DOCX", data=buf.getvalue(),
                    file_name=f"resume_{selected_key}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True)
            except:
                st.info("Install python-docx for DOCX export")

        st.markdown("#### Live Preview")
        st.components.v1.html(html, height=800, scrolling=True)
    else:
        st.info("Fill in your profile to see a live preview of your resume.")
