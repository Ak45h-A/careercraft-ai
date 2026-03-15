import streamlit as st
import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from utils.ai_helper import call_ai, profile_to_text

def render():
    st.markdown('<h1>📄 ATS Resume Builder</h1>', unsafe_allow_html=True)
    st.markdown('<p class="subtitle">Paste a job description → get a tailored, ATS-optimized resume with keyword matching</p>', unsafe_allow_html=True)

    p = st.session_state.profile

    if not p.get("name"):
        st.markdown("""<div class="tip-box">⚠️ Please fill in <strong>My Profile</strong> first before generating a resume.</div>""", unsafe_allow_html=True)
        if st.button("→ Go to My Profile"):
            st.session_state.page = "profile"
            st.rerun()
        return

    tab1, tab2, tab3 = st.tabs(["🎯 Generate Resume", "📊 ATS Analyzer", "📥 Export"])

    with tab1:
        col1, col2 = st.columns([1, 1])

        with col1:
            st.markdown('<div class="step-header">Job Details</div>', unsafe_allow_html=True)
            job_title = st.text_input("Target Job Title *", placeholder="e.g. Senior Python Developer")
            company_name = st.text_input("Company Name", placeholder="e.g. Google, TCS, Infosys")
            job_description = st.text_area("Paste Job Description *",
                placeholder="Paste the full job description here...\n\nThe AI will extract keywords and tailor your resume accordingly.",
                height=250)
            tone = st.selectbox("Resume Tone", ["Professional & Formal", "Technical & Detailed", "Creative & Dynamic", "Executive & Leadership"])
            include_photo = st.checkbox("Include Photo Placeholder", value=False)
            include_objective = st.checkbox("Include Career Objective", value=True)

        with col2:
            st.markdown('<div class="step-header">Profile Preview</div>', unsafe_allow_html=True)
            st.markdown(f"""
            <div class="feature-card">
                <h3>👤 {p.get('name', '—')}</h3>
                <p>📧 {p.get('email','—')} | 📞 {p.get('phone','—')}</p>
                <p>📍 {p.get('city','')}, {p.get('state','')}</p>
                <p>💼 {len(p.get('experiences',[]))} job(s) | 🎓 {len(p.get('education',[]))} degree(s)</p>
                <p>🛠️ {p.get('technical_skills','No skills added')[:80]}...</p>
            </div>
            """, unsafe_allow_html=True)

            st.markdown("""<div class="tip-box">
            <strong>ATS Tips:</strong><br>
            ✅ Use exact keywords from JD<br>
            ✅ Quantify achievements<br>
            ✅ Use standard section headings<br>
            ✅ Avoid tables & graphics in ATS version
            </div>""", unsafe_allow_html=True)

        st.markdown("<hr class='divider'>", unsafe_allow_html=True)
        if st.button("🚀 Generate ATS Resume", use_container_width=True):
            if not job_title or not job_description:
                st.error("Please fill in Job Title and Job Description.")
            elif not st.session_state.api_key:
                st.error("Please add your API key in Settings first.")
            else:
                with st.spinner("🤖 AI is crafting your tailored resume..."):
                    profile_text = profile_to_text(p)
                    prompt = f"""
You are an expert ATS resume writer. Create a highly optimized, ATS-friendly resume.

TARGET ROLE: {job_title}
COMPANY: {company_name}
TONE: {tone}

JOB DESCRIPTION:
{job_description}

CANDIDATE PROFILE:
{profile_text}

INSTRUCTIONS:
1. Analyze the JD and extract key skills, keywords, technologies
2. Tailor the resume to match those keywords naturally
3. Use strong action verbs and quantify achievements
4. Format with clear sections: {"Career Objective, " if include_objective else ""}Contact Info, Professional Summary, Skills, Work Experience, Education, Projects, Certifications, Achievements
5. ATS-optimized: no tables, no graphics, clean formatting
6. Include a "KEYWORD MATCH SECTION" at the end listing matched keywords

Output the resume in clean markdown format with clear section headers.
"""
                    result = call_ai(prompt, max_tokens=3000)
                    st.session_state.generated_resume = result
                    st.success("✅ Resume generated!")

        if st.session_state.generated_resume:
            st.markdown("<hr class='divider'>", unsafe_allow_html=True)
            st.markdown('<div class="step-header">Generated Resume</div>', unsafe_allow_html=True)
            st.markdown(st.session_state.generated_resume)

    with tab2:
        st.markdown('<div class="step-header">ATS Score Analyzer</div>', unsafe_allow_html=True)
        st.markdown("""<div class="tip-box">Paste your existing resume and a job description to get an ATS compatibility score and improvement tips.</div>""", unsafe_allow_html=True)

        col1, col2 = st.columns(2)
        with col1:
            resume_text = st.text_area("Paste Your Resume Text", height=300, placeholder="Paste your current resume text here...")
        with col2:
            jd_text = st.text_area("Paste Job Description", height=300, placeholder="Paste the target job description...")

        if st.button("📊 Analyze ATS Score", use_container_width=True):
            if resume_text and jd_text:
                with st.spinner("Analyzing..."):
                    prompt = f"""
Analyze this resume against the job description for ATS compatibility.

RESUME:
{resume_text}

JOB DESCRIPTION:
{jd_text}

Provide:
1. ATS Score (0-100) with reasoning
2. Matched Keywords (list them)
3. Missing Keywords (list them)
4. Top 5 Improvement Suggestions
5. Section-wise analysis

Format as structured markdown.
"""
                    result = call_ai(prompt)
                    st.markdown(result)
            else:
                st.error("Please paste both resume and job description.")

    with tab3:
        st.markdown('<div class="step-header">Export Your Resume</div>', unsafe_allow_html=True)

        if not st.session_state.generated_resume:
            st.info("Generate a resume first in the 'Generate Resume' tab.")
            return

        st.markdown("""<div class="success-box">✅ Resume ready for export!</div>""", unsafe_allow_html=True)

        col1, col2, col3 = st.columns(3)
        with col1:
            # Export as markdown/txt
            st.download_button(
                label="📥 Download as .txt",
                data=st.session_state.generated_resume,
                file_name="resume.txt",
                mime="text/plain",
                use_container_width=True
            )
        with col2:
            # Export as HTML
            html_content = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<title>Resume - {p.get('name','')}</title>
<style>
  body {{ font-family: 'Segoe UI', sans-serif; max-width: 900px; margin: 40px auto; padding: 20px; color: #333; line-height: 1.6; }}
  h1 {{ color: #2c3e50; font-size: 2rem; margin-bottom: 4px; }}
  h2 {{ color: #6c63ff; font-size: 1.1rem; border-bottom: 2px solid #6c63ff; padding-bottom: 4px; margin-top: 24px; }}
  h3 {{ color: #2c3e50; }}
  ul {{ padding-left: 20px; }}
  li {{ margin-bottom: 4px; }}
  .contact {{ color: #666; font-size: 0.95rem; }}
  @media print {{ body {{ margin: 20px; }} }}
</style>
</head>
<body>
{st.session_state.generated_resume.replace(chr(10), '<br>').replace('# ', '<h1>').replace('## ', '<h2>').replace('### ', '<h3>')}
</body>
</html>"""
            st.download_button(
                label="🌐 Download as HTML",
                data=html_content,
                file_name="resume.html",
                mime="text/html",
                use_container_width=True
            )
        with col3:
            try:
                from docx import Document
                from io import BytesIO
                doc = Document()
                doc.add_heading(p.get("name", "Resume"), 0)
                for line in st.session_state.generated_resume.split("\n"):
                    if line.startswith("## "):
                        doc.add_heading(line[3:], level=1)
                    elif line.startswith("### "):
                        doc.add_heading(line[4:], level=2)
                    elif line.startswith("# "):
                        doc.add_heading(line[2:], level=0)
                    elif line.strip():
                        doc.add_paragraph(line)
                buf = BytesIO()
                doc.save(buf)
                st.download_button(
                    label="📄 Download as DOCX",
                    data=buf.getvalue(),
                    file_name="resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except ImportError:
                st.warning("python-docx not installed. Run: pip install python-docx")

        st.markdown("<hr class='divider'>", unsafe_allow_html=True)
        st.markdown("**PDF Export:** Open the HTML file in Chrome → Press Ctrl+P → Save as PDF for best results.")
