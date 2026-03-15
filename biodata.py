import streamlit as st
import sys, os
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from utils.ai_helper import call_ai

def generate_biodata_html(b: dict, style: str) -> str:
    styles = {
        "Traditional": {
            "bg": "#fff8f0", "header_bg": "#8B1A1A", "accent": "#8B1A1A",
            "font": "Georgia", "border": "#c0392b"
        },
        "Modern Minimal": {
            "bg": "#ffffff", "header_bg": "#2c3e50", "accent": "#6c63ff",
            "font": "Segoe UI", "border": "#6c63ff"
        },
        "Floral Pink": {
            "bg": "#fff0f5", "header_bg": "#d63384", "accent": "#d63384",
            "font": "Palatino", "border": "#f48fb1"
        },
        "Royal Gold": {
            "bg": "#fffef0", "header_bg": "#7d6608", "accent": "#b7860b",
            "font": "Times New Roman", "border": "#f0c040"
        }
    }
    s = styles.get(style, styles["Traditional"])

    def row(label, value):
        return f'<tr><td class="label">{label}</td><td class="value">{value or "—"}</td></tr>' if value else ""

    personal_rows = "".join([
        row("Full Name", b.get("name")),
        row("Date of Birth", b.get("dob")),
        row("Time of Birth", b.get("tob")),
        row("Place of Birth", b.get("pob")),
        row("Age", b.get("age")),
        row("Height", b.get("height")),
        row("Complexion", b.get("complexion")),
        row("Blood Group", b.get("blood_group")),
        row("Religion", b.get("religion")),
        row("Caste / Sub-caste", b.get("caste")),
        row("Gotra", b.get("gotra")),
        row("Nakshatra / Rashi", b.get("rashi")),
        row("Manglik", b.get("manglik")),
    ])

    professional_rows = "".join([
        row("Education", b.get("education")),
        row("Occupation", b.get("occupation")),
        row("Organization", b.get("organization")),
        row("Annual Income", b.get("income")),
        row("City", b.get("city")),
    ])

    family_rows = "".join([
        row("Father's Name", b.get("father")),
        row("Father's Occupation", b.get("father_occ")),
        row("Mother's Name", b.get("mother")),
        row("Mother's Occupation", b.get("mother_occ")),
        row("Brothers", b.get("brothers")),
        row("Sisters", b.get("sisters")),
        row("Family Type", b.get("family_type")),
        row("Native Place", b.get("native")),
    ])

    contact_rows = "".join([
        row("Contact No.", b.get("phone")),
        row("Email", b.get("email")),
        row("Address", b.get("address")),
    ])

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>Biodata - {b.get('name','')}</title>
<style>
  @import url('https://fonts.googleapis.com/css2?family=Noto+Serif:wght@400;700&display=swap');
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{
    font-family: '{s['font']}', 'Noto Serif', serif;
    background: #f0f0f0;
    padding: 20px;
  }}
  .page {{
    max-width: 800px;
    margin: 0 auto;
    background: {s['bg']};
    border: 3px solid {s['border']};
    border-radius: 8px;
    overflow: hidden;
  }}
  .header {{
    background: {s['header_bg']};
    color: white;
    text-align: center;
    padding: 32px 20px;
  }}
  .header .om {{ font-size: 2rem; margin-bottom: 8px; }}
  .header h1 {{ font-size: 1.8rem; font-weight: 700; letter-spacing: 2px; }}
  .header .subtitle {{ font-size: 0.9rem; opacity: 0.85; margin-top: 6px; letter-spacing: 3px; text-transform: uppercase; }}
  .photo-section {{
    display: flex;
    align-items: center;
    gap: 20px;
    padding: 20px 32px;
    border-bottom: 1px solid {s['border']}44;
  }}
  .photo-box {{
    width: 120px;
    height: 150px;
    border: 2px solid {s['border']};
    border-radius: 6px;
    display: flex;
    align-items: center;
    justify-content: center;
    background: #f9f9f9;
    color: #aaa;
    font-size: 0.8rem;
    text-align: center;
    flex-shrink: 0;
  }}
  .name-big {{
    font-size: 1.6rem;
    font-weight: 700;
    color: {s['header_bg']};
  }}
  .tagline {{ font-size: 0.9rem; color: #666; margin-top: 4px; }}
  .body {{ padding: 20px 32px 32px; }}
  .section {{ margin-bottom: 24px; }}
  .section-title {{
    font-size: 0.75rem;
    font-weight: 700;
    text-transform: uppercase;
    letter-spacing: 3px;
    color: {s['accent']};
    border-bottom: 2px solid {s['accent']};
    padding-bottom: 6px;
    margin-bottom: 12px;
  }}
  table {{ width: 100%; border-collapse: collapse; }}
  .label {{
    font-size: 0.85rem;
    color: #666;
    font-weight: 600;
    padding: 6px 0;
    width: 40%;
    vertical-align: top;
  }}
  .value {{
    font-size: 0.9rem;
    color: #222;
    padding: 6px 0 6px 12px;
    border-left: 2px solid {s['border']}33;
  }}
  tr:nth-child(even) td {{ background: {s['accent']}08; }}
  .about {{
    background: {s['accent']}10;
    border-left: 4px solid {s['accent']};
    padding: 12px 16px;
    border-radius: 0 8px 8px 0;
    font-size: 0.9rem;
    color: #444;
    line-height: 1.7;
  }}
  .footer {{
    text-align: center;
    padding: 16px;
    background: {s['header_bg']}15;
    border-top: 1px solid {s['border']}44;
    color: #888;
    font-size: 0.8rem;
    letter-spacing: 1px;
  }}
  @media print {{ body {{ background: white; padding: 0; }} .page {{ border: none; }} }}
</style>
</head>
<body>
<div class="page">
  <div class="header">
    <div class="om">🕉️</div>
    <h1>BIODATA</h1>
    <div class="subtitle">For Marriage Purpose</div>
  </div>

  <div class="photo-section">
    <div class="photo-box">📷<br>Photo<br>Here</div>
    <div>
      <div class="name-big">{b.get('name','Your Name')}</div>
      <div class="tagline">{b.get('education','')} {'| ' + b.get('occupation','') if b.get('occupation') else ''}</div>
      <div style="margin-top:10px; font-size:0.85rem; color:#555;">
        📍 {b.get('city','')}&nbsp;&nbsp;
        📞 {b.get('phone','')}&nbsp;&nbsp;
        ✉️ {b.get('email','')}
      </div>
    </div>
  </div>

  <div class="body">
    {'<div class="section"><div class="section-title">About Me</div><div class="about">' + b.get('about','') + '</div></div>' if b.get('about') else ''}

    {'<div class="section"><div class="section-title">Personal Details</div><table>' + personal_rows + '</table></div>' if personal_rows else ''}

    {'<div class="section"><div class="section-title">Education & Professional</div><table>' + professional_rows + '</table></div>' if professional_rows else ''}

    {'<div class="section"><div class="section-title">Family Details</div><table>' + family_rows + '</table></div>' if family_rows else ''}

    {'<div class="section"><div class="section-title">Partner Preference</div><div class="about">' + b.get('preference','') + '</div></div>' if b.get('preference') else ''}

    {'<div class="section"><div class="section-title">Contact Information</div><table>' + contact_rows + '</table></div>' if contact_rows else ''}
  </div>

  <div class="footer">
    ॐ शुभ विवाह • All information provided is true to the best of our knowledge
  </div>
</div>
</body>
</html>"""


def render():
    st.markdown('<h1>💒 Marriage Biodata Creator</h1>', unsafe_allow_html=True)
    st.markdown('<p class="subtitle">Generate a beautiful, traditional or modern marriage biodata with multiple styles</p>', unsafe_allow_html=True)

    if "biodata" not in st.session_state:
        st.session_state.biodata = {}
    b = st.session_state.biodata

    tab1, tab2 = st.tabs(["📝 Fill Details", "🎨 Preview & Export"])

    with tab1:
        col1, col2 = st.columns(2)

        with col1:
            st.markdown('<div class="step-header">Personal Details</div>', unsafe_allow_html=True)
            b["name"] = st.text_input("Full Name *", value=b.get("name",""), placeholder="Priya Sharma")
            b["dob"] = st.text_input("Date of Birth", value=b.get("dob",""), placeholder="15 March 1997")
            b["tob"] = st.text_input("Time of Birth", value=b.get("tob",""), placeholder="6:30 AM")
            b["pob"] = st.text_input("Place of Birth", value=b.get("pob",""), placeholder="Mumbai, Maharashtra")
            b["age"] = st.text_input("Age", value=b.get("age",""), placeholder="27 years")
            b["height"] = st.text_input("Height", value=b.get("height",""), placeholder="5'4\"")
            b["complexion"] = st.text_input("Complexion", value=b.get("complexion",""), placeholder="Fair / Wheatish")
            b["blood_group"] = st.text_input("Blood Group", value=b.get("blood_group",""), placeholder="B+")

            st.markdown('<div class="step-header" style="margin-top:16px;">Religion & Horoscope</div>', unsafe_allow_html=True)
            b["religion"] = st.text_input("Religion", value=b.get("religion",""), placeholder="Hindu")
            b["caste"] = st.text_input("Caste / Sub-caste", value=b.get("caste",""), placeholder="Brahmin / Iyer")
            b["gotra"] = st.text_input("Gotra", value=b.get("gotra",""), placeholder="Kashyap")
            b["rashi"] = st.text_input("Nakshatra / Rashi", value=b.get("rashi",""), placeholder="Rohini / Vrishabha")
            b["manglik"] = st.selectbox("Manglik", ["", "No", "Yes", "Partial/Anshik"],
                index=["","No","Yes","Partial/Anshik"].index(b.get("manglik","")))

        with col2:
            st.markdown('<div class="step-header">Education & Profession</div>', unsafe_allow_html=True)
            b["education"] = st.text_input("Education", value=b.get("education",""), placeholder="B.E. Computer Science, VTU")
            b["occupation"] = st.text_input("Occupation", value=b.get("occupation",""), placeholder="Software Engineer")
            b["organization"] = st.text_input("Organization", value=b.get("organization",""), placeholder="Infosys, Bengaluru")
            b["income"] = st.text_input("Annual Income", value=b.get("income",""), placeholder="8 LPA")
            b["city"] = st.text_input("Current City", value=b.get("city",""), placeholder="Bengaluru")

            st.markdown('<div class="step-header" style="margin-top:16px;">Family Details</div>', unsafe_allow_html=True)
            b["father"] = st.text_input("Father's Name", value=b.get("father",""), placeholder="Ramesh Sharma")
            b["father_occ"] = st.text_input("Father's Occupation", value=b.get("father_occ",""), placeholder="Retired Bank Manager")
            b["mother"] = st.text_input("Mother's Name", value=b.get("mother",""), placeholder="Sunita Sharma")
            b["mother_occ"] = st.text_input("Mother's Occupation", value=b.get("mother_occ",""), placeholder="Homemaker")
            b["brothers"] = st.text_input("Brothers", value=b.get("brothers",""), placeholder="1 (Married) / None")
            b["sisters"] = st.text_input("Sisters", value=b.get("sisters",""), placeholder="1 (Younger) / None")
            b["family_type"] = st.selectbox("Family Type", ["", "Nuclear", "Joint", "Extended"],
                index=["","Nuclear","Joint","Extended"].index(b.get("family_type","")))
            b["native"] = st.text_input("Native Place", value=b.get("native",""), placeholder="Mysuru, Karnataka")

        st.markdown("<hr class='divider'>", unsafe_allow_html=True)
        col1, col2 = st.columns(2)
        with col1:
            b["phone"] = st.text_input("Contact Number", value=b.get("phone",""), placeholder="+91 9876543210")
            b["email"] = st.text_input("Email", value=b.get("email",""), placeholder="priya@email.com")
            b["address"] = st.text_input("Address", value=b.get("address",""), placeholder="123, MG Road, Bengaluru - 560001")
        with col2:
            b["about"] = st.text_area("About Me (self-introduction)",
                value=b.get("about",""), height=80,
                placeholder="A brief 3-4 line description about your personality, interests, values...")
            b["preference"] = st.text_area("Partner Preference",
                value=b.get("preference",""), height=80,
                placeholder="Looking for an educated, caring partner from a good family...")

        st.markdown("<hr class='divider'>", unsafe_allow_html=True)
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🤖 AI Generate About Me", use_container_width=True):
                with st.spinner("Generating..."):
                    result = call_ai(f"""
Write a warm, genuine 3-4 line 'About Me' section for a marriage biodata.
Name: {b.get('name','')}
Occupation: {b.get('occupation','')}
Education: {b.get('education','')}
City: {b.get('city','')}
Keep it modest, warm, and suitable for a traditional Indian marriage biodata.
Just the text, no extra formatting.
""", max_tokens=200)
                    b["about"] = result
                    st.session_state.biodata = b
                    st.rerun()
        with col2:
            if st.button("💾 Save Biodata", use_container_width=True):
                st.session_state.biodata = b
                st.success("✅ Biodata saved!")

    with tab2:
        st.markdown('<div class="step-header">Preview & Export</div>', unsafe_allow_html=True)

        style = st.selectbox("Choose Style", ["Traditional", "Modern Minimal", "Floral Pink", "Royal Gold"])

        if b.get("name"):
            html = generate_biodata_html(b, style)

            col1, col2 = st.columns(2)
            with col1:
                st.download_button("🌐 Download HTML (Print to PDF)", data=html,
                    file_name=f"biodata_{b.get('name','').replace(' ','_')}.html",
                    mime="text/html", use_container_width=True)
            with col2:
                try:
                    from docx import Document
                    from io import BytesIO
                    doc = Document()
                    doc.add_heading("BIODATA - FOR MARRIAGE PURPOSE", 0)
                    doc.add_heading(b.get("name",""), 1)
                    doc.add_paragraph(f"📞 {b.get('phone','')} | ✉️ {b.get('email','')}")
                    doc.add_heading("Personal Details", 2)
                    for key in ["dob","tob","pob","age","height","complexion","blood_group","religion","caste","gotra","rashi","manglik"]:
                        if b.get(key):
                            doc.add_paragraph(f"{key.replace('_',' ').title()}: {b[key]}")
                    doc.add_heading("Education & Profession", 2)
                    for key in ["education","occupation","organization","income","city"]:
                        if b.get(key):
                            doc.add_paragraph(f"{key.replace('_',' ').title()}: {b[key]}")
                    doc.add_heading("Family Details", 2)
                    for key in ["father","father_occ","mother","mother_occ","brothers","sisters","family_type","native"]:
                        if b.get(key):
                            doc.add_paragraph(f"{key.replace('_',' ').title()}: {b[key]}")
                    if b.get("about"):
                        doc.add_heading("About Me", 2)
                        doc.add_paragraph(b["about"])
                    if b.get("preference"):
                        doc.add_heading("Partner Preference", 2)
                        doc.add_paragraph(b["preference"])
                    buf = BytesIO()
                    doc.save(buf)
                    st.download_button("📄 Download DOCX", data=buf.getvalue(),
                        file_name=f"biodata_{b.get('name','').replace(' ','_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True)
                except:
                    st.info("Install python-docx for DOCX export")

            st.markdown("#### Live Preview")
            st.components.v1.html(html, height=900, scrolling=True)
        else:
            st.info("Fill in at least your name in the 'Fill Details' tab to see a preview.")
